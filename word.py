'''
    #利用python读取word文档，先读取段落
'''
import os
from docx import Document
import difflib
from tqdm import tqdm
import pandas as pd
import time
import concurrent.futures


def word_find(file_name,name):

    #打开word文档
    document = Document(f'{path}\{file_name}')

    #获取所有段落
    all_paragraphs = document.paragraphs
    #打印看看all_paragraphs是什么东西

    #是列表就开始循环读取
    all_list = []
    for paragraph in all_paragraphs:
        #打印每一个段落的文字
        all_list.append(paragraph.text)



    data = difflib.get_close_matches(name, all_list, cutoff=0.1)

    # index_key = [for ]all_list.index(data) + 1
    print(file_name,data)
    # print(index_key)
    return all_list







def find_table(file_name,name):
    '''
    查找 file_name 内 含有name的表格
    :param file_name: word 的文件名
    :param name: 要找的字段
    :return: file_name
    '''
    document = Document(f'{path}\{file_name}')


    for table in document.tables:
        for i,row in enumerate(table.rows):
            if i==0:
                for cell in row.cells:
                    if cell.text == name:
                        # print(file_name)
                        return file_name
    return 'not find'

def find_all_file(path,name):
    '''
    查找表格内含有 '排放浓度' 的所有文件
    :param path:查找路径
    :return:
    '''

    filelist = []
    for file_name in tqdm(os.listdir(path)):
        ans = find_table(file_name, name)
        if ans !='not find':
            filelist.append(ans)
    return filelist

def find_table_return_table(file_name,name):
    '''
    查找 file_name 内 含有name的表格
    :param file_name: word 的文件名
    :param name: 要找的字段
    :return: table
    '''
    document = Document(f'{path}\{file_name}')

    # 放找到的表格
    tables=[]

    table_last_row_first_value = ""
    flag= False
    columns_num=0
    for table in document.tables:
        if flag == 0:
            for i, row in enumerate(table.rows):
                if i == 0:
                    for cell in row.cells:
                        if cell.text == name:
                            flag = True
                            tables.append(table)

            columns_num = len(row.cells)


        elif flag and columns_num!=len(table.columns):
            return tables
        # 找到
        elif flag and columns_num==len(table.columns):

            tables.append(table)




    return  tables
def word_table_to_excel(tables,output):
    '''
    word 表格转换为 excel
    :param table: 
    :param output: 输出路径
    :return: 
    '''
    table_list=[]
    for table in tables:
        for row in table.rows:
            row_list = []
            for cell in row.cells:
                row_list.append(cell.text)
            table_list.append(row_list)

    df = pd.DataFrame(table_list[1:],columns=table_list[0])
    df.to_excel(output,index=False)
    return df

def run_2(file_name):
    tables = find_table_return_table(file_name, '排放浓度')
    if len(tables) > 0:
        word_table_to_excel(tables, f'G:\proprocess\data\output2\{file_name}.xlsx')


if __name__ =="__main__":



    start = time.perf_counter()
    task = list(os.listdir(path))

    with concurrent.futures.ProcessPoolExecutor() as executor:

        results = executor.map(run_2,task)
        for result in results:
            print(result)
    print("多线程",time.perf_counter()-start)
